# GUI Version

import tkinter as tk
from tkinter import filedialog
import requests as req
from bs4 import BeautifulSoup
from docx import Document

class WebScraperGUI:
    def __init__(self, master):
        self.master = master
        master.title("Web Scraper GUI")

        self.url_label = tk.Label(master, text="Enter the URL:")
        self.url_label.pack()

        self.url_entry = tk.Entry(master, width=50)
        self.url_entry.pack()

        self.scrape_button = tk.Button(master, text="Scrape", command=self.scrape)
        self.scrape_button.pack()

    def scrape(self):
        url = self.url_entry.get()

        html_content = fetch_html(url)

        if html_content:
            title, paragraphs = extract_information(html_content)
            save_to_docx(title, paragraphs)
            tk.messagebox.showinfo("Success", "Information saved to Output.docx")

def fetch_html(url):
    try:
        res = req.get(url)
        res.raise_for_status()
        return res.text
    except req.exceptions.RequestException as e:
        print(f"Error fetching content: {e}")
        return None

def extract_information(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')

    title = soup.title.text if soup.title else "No Title"
    paragraphs = [p.text for p in soup.find_all('p')]

    return title, paragraphs

def save_to_docx(title, paragraphs):
    doc = Document()
    doc.add_heading(title, level=1)

    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)

    doc.save('Output.docx')

if __name__ == "__main__":
    root = tk.Tk()
    app = WebScraperGUI(root)
    root.mainloop()
